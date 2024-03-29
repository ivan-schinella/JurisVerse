from utils_module import (
    get_files,
    map_value,
    verifica_file,
    get_parent_directory,
)
import os
from zip_email_module import send_msg
import streamlit as st
from docx import Document

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import Docx2txtLoader
from langchain_community.vectorstores import FAISS
from langchain_openai.embeddings import OpenAIEmbeddings
from langchain_openai.chat_models import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableLambda, RunnablePassthrough
# from langchain_mistralai.chat_models import ChatMistralAI
from langchain.prompts import PromptTemplate
# from langchain_community.embeddings import HuggingFaceEmbeddings
# from langchain_mistralai import MistralAIEmbeddings
from langchain_core.output_parsers import StrOutputParser
# from langchain_anthropic import ChatAnthropic

openai_api_key = os.environ['OPENAI_API_KEY']
# mistral_api_key = os.environ['MISTRAL-API-KEY']
# anthropic_api_key = os.environ['ANTHROPIC-API-KEY']

# embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
# embeddings = HuggingFaceEmbeddings(model_name="paraphrase-multilingual-mpnet-base-v2")
embeddings = OpenAIEmbeddings(api_key=openai_api_key, model="text-embedding-3-small")
# embeddings = MistralAIEmbeddings()

# model = ChatOpenAI(model="gpt-4-turbo-preview", temperature=0.1)
model = ChatOpenAI(api_key=openai_api_key, model="gpt-3.5-turbo", temperature=0.1)
# model = ChatAnthropic(temperature=0.1, model_name="claude-3-opus-20240229")
# model = ChatAnthropic(temperature=0.1, model_name="claude-3-haiku-20240307")
# model = ChatAnthropic(temperature=0.1, model_name="claude-3-sonnet-20240229")
# model = ChatMistralAI(model="mistral-small-2402", temperature=0.1)
# model = ChatMistralAI(model="mistral-small-2312", temperature=0.1)
# model = ChatMistralAI(model="mistral-medium-2312", temperature=0.1)

template = """Answer the following questions based only on the provided context:
{context}
Don't repeat the question within your answer.
Provide a detailed answer.

Questions: {question}
"""
prompt = ChatPromptTemplate.from_template(template)

current_dir = get_parent_directory()

path_docs = os.path.join(current_dir, "docs")
path_abstracts = os.path.join(current_dir, "abstracts")

def interroga():
    st.session_state.query = st.empty()
    with st.session_state.query.container():
        with st.form("query_llm", border=True):
            st.subheader(":blue[Interroga i documenti]")
            st.text_input(
                "Digita la prima domanda:",
                value="In base a quale principio giuridico il giudice ha deciso?",
                key="first_question",
            )
            st.text_input(
                "Digita la seconda domanda:",
                value="Quali sono le leggi citate per motivare la decisione?",
                key="second_question",
            )

            st.form_submit_button(
                label="Interroga", type="primary", on_click=start_query
            )


def start_query():
    st.session_state.query.empty()
    if st.session_state.first_question and not st.session_state.second_question:
        my_prompt = st.session_state.first_question
    elif not st.session_state.first_question and st.session_state.second_question:
        my_prompt = st.session_state.second_question
    else:
        my_prompt = (
            "Domanda n.1: "
            + st.session_state.first_question
            + "\n"
            + "Domanda n.2: "
            + st.session_state.second_question
        )

    with st.spinner(text="Creazione lista sentenze..."):
        files = get_files(path_docs)

    num_sentenze = len(files)
    my_index = 0

    if num_sentenze >= 1:

        my_bar = st.progress(0)

        document = Document()
        document.add_heading("Abstracts dei Ricorsi", 0)
        document.add_heading(my_prompt, level=1)

        for file in files:
            my_index += 1

            mapped_value = int(
                map_value(
                    my_index,
                    from_min=0,
                    from_max=num_sentenze,
                    to_min=0,
                    to_max=100,
                )
            )
            my_bar.progress(
                mapped_value,
                text=(
                    f"Elaborazione sentenza {my_index} di {num_sentenze} in corso..."
                ),
            )

            # path_abstracts = os.path.join(current_dir, "abstracts")
            loader = Docx2txtLoader(f'{path_docs}/{file}')
            docs = loader.load()

            text = ""
            for page in docs:
                text += page.page_content
            text = text.replace("\t", " ")

            text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
                chunk_size=2000, chunk_overlap=200
            )
            texts = text_splitter.split_text(text)

            vectorstore = FAISS.from_texts(texts, embedding=embeddings)
            retriever = vectorstore.as_retriever(search_kwargs={"k": 10})

            chain = (
                {"context": retriever, "question": RunnablePassthrough()}
                | prompt
                | model
                | StrOutputParser()
            )

            # answer = chain.invoke(my_prompt)
            answer = "test attivato in modulo [query_module.py]"

            if answer:
                st.session_state.abstract_creati = True

            document.add_heading(file, level=2)
            document.add_paragraph(answer)

        my_bar.empty()
        document.save(f'{path_abstracts}/Abstracts.docx')

        # Invio automatico della email
        if st.session_state.auto_send_email and verifica_file(
            path_abstracts, "Abstracts.docx"
        ):
            if not st.session_state.email_sent:
                with st.spinner("Preparazione ed invio email..."):
                    send_msg()
                    st.session_state.email_sent = True
                st.toast("Email inviata!")
